# app/logic/langchain_helper.py
"""
LangChain / RAG helper for the No-Code Chatbot platform.

Responsibilities:
- Manage DB reads/writes for user conversation history
- Provide `get_bot_response(user_id, message)` as the single entrypoint used by API
- Provide `get_rag_response(...)` which performs retrieval + LLM invocation
- Provide utility functions for indexing / initializing vectorstore (stubs included)
"""

from __future__ import annotations
import os
import json
import logging
import re
import time
from typing import Any, Dict, List, Optional, Tuple

from langchain_groq import ChatGroq
from langchain_huggingface import HuggingFaceEmbeddings 
from langchain_community.vectorstores import FAISS

from langchain_text_splitters import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
from langchain_core.documents import Document
import uuid
import pathlib
from passlib.context import CryptContext
import numpy as np
from collections import Counter

import requests
from typing import BinaryIO
import pdfplumber
import docx  # python-docx
from bs4 import BeautifulSoup
import csv
from io import StringIO

import psycopg2  
from dotenv import load_dotenv

    
load_dotenv()
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)


# ---------------------------
# Config placeholders (FILL IN)
# ---------------------------
DB_CONFIG = {
    # Example keys (fill from env or other config)
    "dbname": os.getenv("POSTGRES_DB"),
    "user": os.getenv("POSTGRES_USER"),
    "password": os.getenv("POSTGRES_PASSWORD"),
    "host": os.getenv("POSTGRES_HOST", "localhost"),
    "port": os.getenv("POSTGRES_PORT", "5432"),
    # Add sslmode or other params if required
}

# LLM / embeddings / vectorstore placeholders: instantiate them in init_llm_embeddings()
LLM = None
EMBEDDINGS = None
CROSS_ENCODER = None
VECSTORE = None
RETRIEVER = None


# ==================== IMPROVED RAG PROMPT TEMPLATE ====================
# Enhanced with better structure and few-shot examples (no citations)

RAG_PROMPT_TEMPLATE = """You are an expert question-answering assistant. Your role is to provide accurate, helpful answers based ONLY on the provided context.

## CRITICAL RULES:
1. Answer ONLY using information explicitly stated in the CONTEXT below
2. If the answer is not in the context, say "I don't have enough information in the provided context to answer this question accurately."
3. Be concise but complete
4. Follow the BOT INSTRUCTIONS precisely

## BOT INSTRUCTIONS:
{bot_instructions}

## CONTEXT:
{retrieved_chunks}

## CONVERSATION HISTORY:
{history}

## USER QUESTION:
{query}

## RESPONSE GUIDELINES:
- If the context contains the answer, provide it clearly
- If the context is insufficient, politely state that you need more information
- Do not make up information or use knowledge outside the provided context
- Format your response naturally and conversationally
- Include source references in parentheses where appropriate based on the [Source: ...] blocks.
- If multiple sources are used, cite each one separately.
- If the answer cannot be clearly supported by the provided context, explicitly say so.

## EXAMPLE RESPONSES:

Example 1 (Answer found):
Question: "What is the return policy?"
Context: "Our return policy allows returns within 30 days of purchase."
Response: "According to our policy, you can return items within 30 days of purchase."

Example 2 (Answer not found):
Question: "What is your shipping cost?"
Context: "We offer free shipping on orders over $50."
Response: "I don't have enough information in the provided context to answer this question accurately. The context mentions free shipping for orders over $50, but doesn't specify the shipping cost for smaller orders."

Now, provide your response to the user's question:"""

# Alternative simpler template (use if the above is too verbose)
RAG_PROMPT_TEMPLATE_SIMPLE = """You are a helpful assistant. Answer the question using ONLY the provided context. If the answer isn't in the context, say so.

BOT INSTRUCTIONS:
{bot_instructions}

CONTEXT:
{retrieved_chunks}

CONVERSATION HISTORY:
{history}

QUESTION: {query}

Answer:"""


pwd_context = CryptContext(
    schemes=["pbkdf2_sha256"],
    deprecated="auto"
)

def hash_password(password: str) -> str:
    return pwd_context.hash(password)


def verify_password(password: str, hashed: str) -> bool:
    return pwd_context.verify(password, hashed)


def create_user(email: str, password: str) -> str:
    conn = _get_db_conn()
    cur = conn.cursor()

    user_id = str(uuid.uuid4())
    cur.execute(
        """
        INSERT INTO users (id, email, password_hash)
        VALUES (%s, %s, %s)
        """,
        (user_id, email, hash_password(password))
    )

    conn.commit()
    cur.close()
    conn.close()
    return user_id


def login_user(email: str, password: str) -> str | None:
    conn = _get_db_conn()
    cur = conn.cursor()

    cur.execute(
        "SELECT id, password_hash FROM users WHERE email=%s",
        (email,)
    )
    row = cur.fetchone()

    cur.close()
    conn.close()

    if not row:
        return None

    user_id, password_hash = row
    return str(user_id) if verify_password(password, password_hash) else None


def create_bot(user_id: str, name: str, instructions: str | None = None) -> str:
    conn = _get_db_conn()
    cur = conn.cursor()

    bot_id = str(uuid.uuid4())
    cur.execute(
        """
        INSERT INTO bots (id, user_id, name, instructions, csv_files, doc_files, csv_content, doc_content)
        VALUES (%s, %s, %s, %s, '[]'::jsonb, '[]'::jsonb, '[]'::jsonb, '[]'::jsonb)
        """,
        (bot_id, user_id, name, instructions)
    )

    conn.commit()
    cur.close()
    conn.close()
    return bot_id


def get_bot(bot_id: str):
    conn = _get_db_conn()
    cur = conn.cursor()

    cur.execute(
        "SELECT id, name, instructions, csv_files, doc_files, csv_content, doc_content FROM bots WHERE id = %s",
        (bot_id,)
    )
    row = cur.fetchone()

    cur.close()
    conn.close()

    if not row:
        return None

    return {
        "id": row[0],
        "name": row[1],
        "instructions": row[2],
        "csv_files": row[3] if row[3] else [],
        "doc_files": row[4] if row[4] else [],
        "csv_content": row[5] if row[5] else [],
        "doc_content": row[6] if row[6] else []
    }


def _get_db_conn():
    """Create a new DB connection using DB_CONFIG. Caller must close the connection."""
    if not DB_CONFIG.get("dbname"):
        raise RuntimeError("DB_CONFIG is not configured. Fill in DB credentials.")
    conn = psycopg2.connect(**DB_CONFIG)
    return conn


def store_message(user_id: str, bot_id: str, role: str, text: str):
    conn = _get_db_conn()
    cur = conn.cursor()

    cur.execute(
        """
        INSERT INTO conversations (id, user_id, bot_id, role, message)
        VALUES (%s, %s, %s, %s, %s)
        """,
        (str(uuid.uuid4()), user_id, bot_id, role, text)
    )

    conn.commit()
    cur.close()
    conn.close()


def get_conversation_history(user_id: str, bot_id: str, limit: int = 6):
    conn = _get_db_conn()
    cur = conn.cursor()

    cur.execute(
        """
        SELECT role, message
        FROM conversations
        WHERE user_id=%s AND bot_id=%s
        ORDER BY created_at DESC
        LIMIT %s
        """,
        (user_id, bot_id, limit)
    )

    rows = cur.fetchall()
    cur.close()
    conn.close()

    return list(reversed(rows))


def delete_faiss_index_for_bot(bot_id: str):
    import shutil
    index_dir = pathlib.Path("data") / f"faiss_index_{bot_id}"
    if index_dir.exists():
        shutil.rmtree(index_dir)


def delete_csv_faiss_for_bot(bot_id: str):
    import shutil
    index_dir = pathlib.Path("data") / f"faiss_index_{bot_id}_csv"
    if index_dir.exists():
        shutil.rmtree(index_dir)


def get_bot_instructions(bot_id: str) -> str:
    conn = _get_db_conn()
    cur = conn.cursor()
    cur.execute("SELECT instructions FROM bots WHERE id = %s", (bot_id,))
    row = cur.fetchone()
    cur.close()
    conn.close()
    return row[0] if row and row[0] else ""


def init_llm_embeddings_and_retriever():
    global LLM, EMBEDDINGS, CROSS_ENCODER, VECSTORE, RETRIEVER

    if LLM is not None and RETRIEVER is not None:
        return  # already initialized

    # ------------------------
    # 1. Load LLM (Groq)
    # ------------------------
    groq_api_key = os.getenv("GROQ_API_KEY")
    if not groq_api_key:
        raise ValueError("Missing GROQ_API_KEY in .env")

    LLM = ChatGroq(
        groq_api_key=groq_api_key,
        model_name="llama-3.1-8b-instant",   # <=== Your chosen model
        temperature=0.2             # low temp for RAG accuracy
    )

    # ------------------------
    # 2. Load Embeddings
    # IMPROVED: Configurable embedding model
    # ------------------------
    embedding_model = os.getenv("EMBEDDING_MODEL", "sentence-transformers/all-mpnet-base-v2")
    # Options:
    # - "sentence-transformers/all-mpnet-base-v2" (default, 768 dims, good balance)
    # - "BAAI/bge-large-en-v1.5" (1024 dims, better accuracy, slower)
    # - "intfloat/e5-large-v2" (1024 dims, excellent for retrieval)
    # - "sentence-transformers/all-MiniLM-L6-v2" (384 dims, faster, good for large datasets)
    
    EMBEDDINGS = HuggingFaceEmbeddings(
        model_name=embedding_model,
        model_kwargs={'device': 'cpu'},  # Use 'cuda' if GPU available
        encode_kwargs={'normalize_embeddings': True}  # Normalize for better cosine similarity
    )

    # ------------------------
    # 3. Load Cross-Encoder
    # ------------------------
    try:
        from sentence_transformers import CrossEncoder
        CROSS_ENCODER = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2", device='cpu')
    except ImportError:
        logger.warning("sentence_transformers not installed. Cross-encoder reranking will be disabled.")
        CROSS_ENCODER = None

    # For local testing we rely ONLY on per-bot FAISS indexes
    VECSTORE = None
    RETRIEVER = None


def expand_query(query: str, llm: Any, max_expansions: int = 3) -> List[str]:
    """
    Generate query expansions using LLM to improve retrieval.
    Returns list of query variations including the original.
    """
    if llm is None:
        return [query]
    
    try:
        expansion_prompt = f"""Generate {max_expansions} alternative phrasings or expansions of this query that a user might use to find the same information. Return only the queries, one per line, without numbering.

Query: {query}

Alternative queries:"""
        
        response = llm.invoke(expansion_prompt) if hasattr(llm, "invoke") else llm(expansion_prompt)
        expansions = []
        
        if hasattr(response, "content"):
            text = response.content
        else:
            text = str(response)
        
        # Parse expansions
        for line in text.strip().split("\n"):
            line = line.strip()
            if line and not line.startswith("#") and len(line) > 5:
                # Remove numbering if present
                line = re.sub(r'^\d+[\.\)]\s*', '', line)
                expansions.append(line)
        
        # Always include original query first
        result = [query]
        result.extend([e for e in expansions[:max_expansions] if e.lower() != query.lower()])
        return result[:max_expansions + 1]
    except Exception as e:
        logger.warning(f"Query expansion failed: {e}")
        return [query]


def contextualize_query(query: str, history_pairs: List[Tuple[str, str]], llm: Any) -> str:
    """
    Rewrite the user's query based on conversation history to resolve pronouns
    and missing entities. 
    
    CONSTRAINTS: Do not answer the question. Do not invent entities. Only expand 
    missing references. Return ONLY the rewritten query.
    """
    if not history_pairs or not llm:
        return query

    history_str = "\n".join([f"{'User' if s == 'user' else 'Assistant'}: {m}" for s, m in history_pairs])
    
    prompt = f"""Given a chat history and the latest user question which might reference context in the chat history, formulate a standalone question which can be understood without the chat history.

CRITICAL CONSTRAINTS:
1. Do NOT answer the question.
2. Do NOT invent new entities.
3. ONLY expand missing references (like "it", "they", "this").
4. Preserve the original meaning.
5. Return ONLY the rewritten standalone query text, nothing else.

Chat History:
{history_str}

Latest Question: {query}

Standalone Question:"""

    try:
        response = llm.invoke(prompt) if hasattr(llm, "invoke") else llm(prompt)
        result = response.content if hasattr(response, "content") else str(response)
        result = result.strip()
        
        # Fallback if the LLM hallucinated a conversational prefix
        invalid_prefixes = ["standalone question:", "here is the", "the standalone"]
        if any(result.lower().startswith(p) for p in invalid_prefixes):
            return query # Abort if LLM failed constraint

        return result
    except Exception as e:
        logger.warning(f"Contextualization failed: {e}")
        return query


def cross_encoder_rerank(query: str, documents: List[Document], top_k: int = 5) -> List[Document]:
    """
    Rerank documents using a true cross-encoder model that scores (query, document) pairs directly.
    """
    global CROSS_ENCODER
    if not documents:
        return []
    
    if CROSS_ENCODER is None:
        logger.warning("CROSS_ENCODER not initialized. Returning original top_k.")
        return documents[:top_k]
    
    try:
        # Prepare pairs for cross-encoder
        pairs = [[query, doc.page_content] for doc in documents]
        
        # Predict relevance scores
        scores = CROSS_ENCODER.predict(pairs)
        
        # Log top 10 candidate scores
        logger.info("\n[Cross-Encoder] Top 10 Candidate Scores:")
        sorted_indices = np.argsort(scores)[::-1]
        for idx in sorted_indices[:10]:
            logger.info(f"Score: {scores[idx]:.4f} | Source: {documents[idx].metadata.get('source', 'unknown')}")
            
        # Select top k
        top_k_indices = sorted_indices[:top_k]
        
        # Return documents unmodified 
        return [documents[i] for i in top_k_indices]

    except Exception as e:
        logger.exception(f"Cross-encoder reranking failed: {e}")
        return documents[:top_k]


def mmr_selection(query_embedding: List[float], doc_embeddings: List[List[float]], documents: List[Document], top_k: int = 5, lambda_mult: float = 0.7) -> List[Document]:
    """
    Maximal Marginal Relevance (MMR) selection to reduce redundancy among retrieved chunks.
    Selects documents that are relevant to the query but diverse from each other.
    """
    if not documents:
        return []

    # If asking for more/equal than the pool, return all
    if top_k >= len(documents):
        return documents

    try:
        from sklearn.metrics.pairwise import cosine_similarity
        
        # Compute query-doc similarities
        q_doc_sims = cosine_similarity([query_embedding], doc_embeddings)[0]
        
        # Compute doc-doc similarities
        doc_doc_sims = cosine_similarity(doc_embeddings, doc_embeddings)
        
        selected_indices = []
        unselected_indices = list(range(len(documents)))
        
        # 1. Select highest scoring doc first
        best_idx = int(np.argmax(q_doc_sims))
        selected_indices.append(best_idx)
        unselected_indices.remove(best_idx)
        
        # 2. Iteratively select next doc maximizing: \lambda * relevance - (1 - \lambda) * similarity_to_selected
        while len(selected_indices) < top_k and unselected_indices:
            best_score = -float('inf')
            best_idx = -1
            
            for unsel_idx in unselected_indices:
                relevance = q_doc_sims[unsel_idx]
                
                # Max similarity to already selected docs
                max_sim_to_selected = max([doc_doc_sims[unsel_idx][sel_idx] for sel_idx in selected_indices])
                
                # MMR Score
                mmr_score = (lambda_mult * relevance) - ((1 - lambda_mult) * max_sim_to_selected)
                
                if mmr_score > best_score:
                    best_score = mmr_score
                    best_idx = unsel_idx
                    
            selected_indices.append(best_idx)
            unselected_indices.remove(best_idx)
            
        return [documents[i] for i in selected_indices]
        
    except Exception as e:
        logger.warning(f"MMR selection failed: {e}. Falling back to default ordering.")
        return documents[:top_k]


def format_context(documents: List[Document]) -> str:
    """
    Format retrieved documents as plain text for the prompt, explicitly injecting
    metadata [Source: X | Page: Y] delimiters above each chunk.
    """
    if not documents:
        return "No relevant document context found."
    
    formatted_chunks = []
    for doc in documents:
        # Extract metadata safely
        source = doc.metadata.get("source", "Unknown Document")
        page = doc.metadata.get("page")
        
        if page is not None and str(page).strip():
            meta_header = f"[Source: {source} | Page: {page}]"
        else:
            meta_header = f"[Source: {source}]"
            
        chunk_text = f"{meta_header}\n{doc.page_content.strip()}"
        formatted_chunks.append(chunk_text)
        
    return "\n\n====================\n\n".join(formatted_chunks)


def get_rag_response(
    message: str,
    history_pairs: List[Tuple[str, str]],
    retriever: Any,
    llm: Any,
    bot_id: str,
    prompt_template: Optional[str] = None,
    use_query_expansion: bool = True,
    use_reranking: bool = True,
    ) -> Dict[str, Any]:

    if retriever is None or llm is None:
        raise RuntimeError("Retriever and LLM must be provided to run RAG. Call init_llm_embeddings_and_retriever() first.")

    prompt_template = prompt_template or RAG_PROMPT_TEMPLATE
    history_str = "\n".join([f"{'User' if s == 'user' else 'Assistant'}: {m}" for s, m in history_pairs]) if history_pairs else "No previous conversation."

    # Step 1: Query expansion (optional)
    queries_to_try = [message]
    if use_query_expansion:
        try:
            expanded = expand_query(message, llm, max_expansions=2)
            queries_to_try = expanded
            logger.info(f"Query expansion: {queries_to_try}")
        except Exception as e:
            logger.warning(f"Query expansion failed: {e}")

    # Step 2: Retrieve relevant docs (try multiple queries if expanded)
    all_docs = []
    seen_content = set()
    
    for query_variant in queries_to_try:
        try:
            if isinstance(retriever, list):
                docs = retriever
            elif hasattr(retriever, "invoke"):
                docs = retriever.invoke(query_variant)
            else:
                docs = retriever.get_relevant_documents(query_variant)
            
            # Deduplicate by content
            for doc in docs:
                content_hash = hash(doc.page_content[:100])  # Use first 100 chars as hash
                if content_hash not in seen_content:
                    seen_content.add(content_hash)
                    all_docs.append(doc)
        except Exception as e:
            logger.exception("Retriever invocation failed: %s", e)
    
    # Step 3: Rerank documents (optional)
    if use_reranking and all_docs:
        try:
            global EMBEDDINGS
            if EMBEDDINGS is not None:
                all_docs = rerank_documents(message, all_docs, EMBEDDINGS, top_k=min(10, len(all_docs)))
        except Exception as e:
            logger.warning(f"Reranking failed: {e}")

    # Step 4: Format context (plain text, no citations)
    retrieved_chunks = format_context(all_docs)
    bot_instructions = get_bot_instructions(bot_id)

    prompt_text = prompt_template.format(
        bot_instructions=bot_instructions,
        retrieved_chunks=retrieved_chunks,
        history=history_str,
        query=message
    )

    print("\n================ PROMPT SENT TO LLM ================\n")
    print(prompt_text[:6000])   # limit to avoid terminal overflow
    print("\n====================================================\n")


    # Invoke the LLM with the prepared prompt.
    try:
        raw_result = llm.invoke(prompt_text) if hasattr(llm, "invoke") else llm(prompt_text)
    except Exception as e:
        logger.exception("LLM invocation failed: %s", e)
        return {"action": "reply", "text": "Sorry, I had trouble processing that. Please try again later."}

    # The project previously expected strict JSON or a fenced ```json block; normalize that:
    if hasattr(raw_result, "content"):
        result_str = raw_result.content
    else:
        result_str = str(raw_result)

    # Normalize line endings and keep formatting
    result_str = result_str.replace("\r\n", "\n").strip()

    #print("bot:", result_str)

    # Remove triple-backtick fenced json blocks if present
    if result_str.startswith("```json"):
        result_str = re.sub(r"^```json", "", result_str)
        result_str = re.sub(r"```$", "", result_str).strip()

    try:
        parsed = json.loads(result_str)
        return parsed
    except json.JSONDecodeError:
    # This is NORMAL for conversational responses
        return {
            "type": "text",
            "content": result_str
        }


def retrieve_csv_with_scores(bot_id: str, query: str, k: int = 5):
    """
    Retrieve CSV docs with similarity scores.
    Returns list of (Document, score)
    """
    global EMBEDDINGS
    if EMBEDDINGS is None:
        init_llm_embeddings_and_retriever()

    index_dir = pathlib.Path("data") / f"faiss_index_{bot_id}_csv"
    if not index_dir.exists():
        return []

    vs = FAISS.load_local(
        str(index_dir),
        EMBEDDINGS,
        allow_dangerous_deserialization=True
    )

    return vs.similarity_search_with_score(query, k=k)


def index_documents_for_bot(bot_id: str, documents: List[Dict[str, Any]], *, chunk_size: int = 1000, chunk_overlap: int = 200) -> Dict[str, Any]:
    """
    Index a list of documents for a specific bot using FAISS + EMBEDDINGS.
    IMPROVED: Better chunking strategy with adaptive sizes and semantic awareness.
    
    - bot_id: unique id for this bot (used to persist per-bot index)
    - documents: list of dicts: {"text": str, "source": str, "page": int, "meta": {...}}
    - chunk_size / chunk_overlap: splitting parameters (characters)
    Returns: {"bot_id":..., "index_path":..., "num_docs_indexed":..., "num_chunks":...}
    """
    # Lazy init of embeddings if not already
    global EMBEDDINGS

    if EMBEDDINGS is None:
        # Try to initialize the stack (this will raise if init isn't implemented)
        try:
            init_llm_embeddings_and_retriever()
        except Exception as e:
            raise RuntimeError("Embeddings not initialized; call init_llm_embeddings_and_retriever() first.") from e

    # IMPROVED: Semantic Chunking Strategy (Paragraph First)
    # Using 600-900 chars based on token recommendations with small overlap
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size if chunk_size else 800,  # Optimal size for context limit
        chunk_overlap=chunk_overlap if chunk_overlap else 100,  # Small overlap
        separators=[
            "\n\n",    # Paragraphs (High Priority)
            "\n",      # Line breaks
            ". ",      # Sentences
            " ",       # Words
            ""         # Characters (last resort)
        ],
        length_function=len,
        is_separator_regex=False,
        keep_separator=True  # Keep separators for better context
    )

    # Prepare lists for texts and metadatas
    texts: List[str] = []
    metadatas: List[Dict[str, Any]] = []

    num_docs = 0
    num_chunks = 0
    total_chunk_length = 0

    for doc in documents:
        num_docs += 1
        full_text = doc.get("text", "")
        source = doc.get("source", doc.get("filename", "unknown"))
        page = doc.get("page", None)
        meta = doc.get("meta", {}) or {}

        if not full_text or not isinstance(full_text, str):
            continue

        chunks = splitter.split_text(full_text)
        doc_chunk_count = 0
        for i, chunk in enumerate(chunks):
            # Skip extremely small chunks to avoid noise
            if len(chunk.strip()) < 50:
                continue
                
            chunk_id = str(uuid.uuid4())
            texts.append(chunk)
            metadatas.append({
                "bot_id": bot_id,
                "source": source,
                "page": page,
                "chunk_index": i,
                "chunk_id": chunk_id,
                "original_meta": meta
            })
            num_chunks += 1
            doc_chunk_count += 1
            total_chunk_length += len(chunk)
            
        logger.info(f"Document {source} split into {doc_chunk_count} chunks.")

    if num_chunks > 0:
        logger.info(f"Average chunk length: {total_chunk_length / num_chunks:.1f} characters")

    # index path per bot
    index_dir = pathlib.Path("data") / f"faiss_index_{bot_id}"
    index_dir.mkdir(parents=True, exist_ok=True)
    index_path = str(index_dir)

    # If there's no chunks, return early
    if not texts:
        return {"bot_id": bot_id, "index_path": index_path, "num_docs_indexed": num_docs, "num_chunks": 0}

    # If a FAISS index already exists for this bot, load and add; otherwise create new
    try:
        # Always create a fresh FAISS index (safe for retrain)
        vs = FAISS.from_texts(texts, EMBEDDINGS, metadatas=metadatas)
        vs.save_local(index_path)

    except Exception as e:
        # attempt a safer recreate: create fresh index and write once
        try:
            vs = FAISS.from_texts(texts, EMBEDDINGS, metadatas=metadatas)
            vs.save_local(index_path)
        except Exception as e2:
            raise RuntimeError(f"Failed to create or update FAISS index for bot {bot_id}: {e2}") from e2

    # Optionally update a global mapping VECSTORE if you want a single in-memory retriever (not required)
    # For per-bot retriever you can instantiate when answering queries by loading this index.

    return {"bot_id": bot_id, "index_path": index_path, "num_docs_indexed": num_docs, "num_chunks": num_chunks}


def update_bot_name(bot_id: str, new_name: str):
    conn = _get_db_conn()
    cur = conn.cursor()
    cur.execute(
        "UPDATE bots SET name=%s WHERE id=%s",
        (new_name, bot_id)
    )
    conn.commit()
    cur.close()
    conn.close()


def is_bot_trained(bot_id: str) -> bool:
    import pathlib

    doc_index = pathlib.Path("data") / f"faiss_index_{bot_id}"
    csv_index = pathlib.Path("data") / f"faiss_index_{bot_id}_csv"

    return doc_index.exists() or csv_index.exists()


def update_bot_instructions(bot_id: str, instructions: str):
    conn = _get_db_conn()
    cur = conn.cursor()
    cur.execute(
        "UPDATE bots SET instructions=%s WHERE id=%s",
        (instructions, bot_id)
    )
    conn.commit()
    cur.close()
    conn.close()


def parse_and_validate_csv(csv_bytes: bytes):
    """
    Returns list of dicts:
    [{ "question": "...", "answer": "..." }, ...]
    Raises ValueError if invalid.
    """
    text = csv_bytes.decode("utf-8", errors="ignore")
    reader = csv.DictReader(StringIO(text))

    required_cols = {"question", "answer"}
    if not required_cols.issubset(reader.fieldnames or []):
        raise ValueError("CSV must contain columns: question, answer")

    rows = []
    for i, row in enumerate(reader):
        q = (row.get("question") or "").strip()
        a = (row.get("answer") or "").strip()
        if not q or not a:
            continue
        rows.append({
            "question": q,
            "answer": a,
            "row": i
        })

    if not rows:
        raise ValueError("CSV contains no valid question-answer rows")

    return rows


def load_retriever_for_bot(bot_id: str, k: int = 5, score_threshold: float = None):
    """
    Load per-bot FAISS index and return a retriever configured with top-k.
    IMPROVED: Added score threshold filtering.
    
    Raises RuntimeError if embeddings or index are not available.
    """
    global EMBEDDINGS
    if EMBEDDINGS is None:
        try:
            init_llm_embeddings_and_retriever()
        except Exception as e:
            raise RuntimeError("Embeddings not initialized.") from e

    index_dir = pathlib.Path("data") / f"faiss_index_{bot_id}"
    if not index_dir.exists():
        raise RuntimeError(f"No index found for bot_id={bot_id} at {index_dir}")

    try:
        vs = FAISS.load_local(str(index_dir), EMBEDDINGS, allow_dangerous_deserialization=True)
    except Exception as e:
        raise RuntimeError(f"Failed to load FAISS index for bot {bot_id}: {e}") from e

    # Use MMR (Maximal Marginal Relevance) for diversity, or similarity for relevance
    search_kwargs = {"k": k}
    if score_threshold:
        search_kwargs["score_threshold"] = score_threshold
    
    retriever = vs.as_retriever(
        search_type="similarity",  # Can also use "mmr" for diversity
        search_kwargs=search_kwargs
    )
    return retriever


def hybrid_search(
    query: str,
    vectorstore: Any,
    k: int = 15
) -> List[Document]:
    """
    Two-Stage Hybrid Search pooling:
    Stage 1: Retrieve top-k using dense vectors (FAISS)
    Stage 2: Retrieve top-k using sparse keyword matching (BM25)
    Merge and deduplicate candidate pools.
    """
    if vectorstore is None:
        return []
    
    try:
        from rank_bm25 import BM25Okapi
    except ImportError:
        logger.warning("rank_bm25 not installed. Falling back to pure FAISS semantic search.")
        return vectorstore.similarity_search(query, k=k)

    try:
        # Stage 1: Get Semantic Docs
        semantic_docs = vectorstore.similarity_search(query, k=k)
        
        # We need ALL documents in the index to run BM25 against the corpus.
        # Since FAISS doesn't have a direct 'get_all' easily out of the box, we
        # search with a massive k to grab the corpus for the BM25 model on the fly.
        # In production, you'd maintain a persistent BM25 index.
        corpus_docs = vectorstore.similarity_search("", k=10000) 
        
        if not corpus_docs:
            return semantic_docs
            
        # Tokenize corpus for BM25
        tokenized_corpus = [doc.page_content.lower().split() for doc in corpus_docs]
        bm25 = BM25Okapi(tokenized_corpus)
        
        # Tokenize query
        tokenized_query = query.lower().split()
        
        # Stage 2: Get BM25 Docs
        bm25_scores = bm25.get_scores(tokenized_query)
        # Get top K indices
        top_n = np.argsort(bm25_scores)[::-1][:k]
        keyword_docs = [corpus_docs[i] for i in top_n if bm25_scores[i] > 0]
        
        # Stage 3: Merge and Deduplicate Candidates
        merged_docs = []
        seen_chunks = set()
        
        # Interleave to ensure a mix of both strategies
        for s_doc, k_doc in zip(semantic_docs + [None]*k, keyword_docs + [None]*k):
            if s_doc:
                hash_id = hash(s_doc.page_content[:100])
                if hash_id not in seen_chunks:
                    seen_chunks.add(hash_id)
                    merged_docs.append(s_doc)
            if k_doc:
                hash_id = hash(k_doc.page_content[:100])
                if hash_id not in seen_chunks:
                    seen_chunks.add(hash_id)
                    merged_docs.append(k_doc)
                    
        return merged_docs[:k]
        
    except Exception as e:
        logger.warning(f"Two-stage hybrid search failed: {e}. Falling back to FAISS.")
        return vectorstore.similarity_search(query, k=k) if vectorstore else []


def index_csv_qa_faiss(bot_id: str, csv_rows: list, replace_content: bool = True):
    global EMBEDDINGS
    if EMBEDDINGS is None:
        init_llm_embeddings_and_retriever()

    # Store CSV rows in DB – always treat input as the full, final set.
    conn = _get_db_conn()
    cur = conn.cursor()
    cur.execute(
        "UPDATE bots SET csv_content = %s::jsonb WHERE id = %s",
        (json.dumps(csv_rows), bot_id)
    )
    conn.commit()
    cur.close()
    conn.close()

    texts = []
    metadatas = []

    for row in csv_rows:
        embedding_text = f"{row['question'].strip()} {row['answer'].strip()}"
    
        texts.append(embedding_text)
    
        metadatas.append({
            "type": "csv",
            "question": row["question"].strip(),
            "answer": row["answer"].strip(),
            "row": row["row"],
            "bot_id": bot_id
        })


    index_dir = pathlib.Path("data") / f"faiss_index_{bot_id}_csv"
    index_dir.mkdir(parents=True, exist_ok=True)

    if index_dir.exists() and any(index_dir.iterdir()):
        # 🔥 LOAD + APPEND
        vs = FAISS.load_local(
            str(index_dir),
            EMBEDDINGS,
            allow_dangerous_deserialization=True
        )
        vs.add_texts(texts, metadatas=metadatas)
    else:
        vs = FAISS.from_texts(texts, EMBEDDINGS, metadatas=metadatas)

    vs.save_local(str(index_dir))


def load_csv_retriever_for_bot(bot_id: str, k: int = 5):
    global EMBEDDINGS
    if EMBEDDINGS is None:
        init_llm_embeddings_and_retriever()

    index_dir = pathlib.Path("data") / f"faiss_index_{bot_id}_csv"
    if not index_dir.exists():
        return None

    vs = FAISS.load_local(
        str(index_dir),
        EMBEDDINGS,
        allow_dangerous_deserialization=True
    )

    return vs.as_retriever(search_kwargs={"k": k})


def delete_bot(bot_id: str):
    """
    Completely delete a bot:
    - conversations
    - vector indexes
    - bot record
    """
    conn = _get_db_conn()
    cur = conn.cursor()

    try:
        # 1️⃣ Delete conversations
        cur.execute(
            "DELETE FROM conversations WHERE bot_id = %s",
            (bot_id,)
        )

        # 2️⃣ Delete bot itself
        cur.execute(
            "DELETE FROM bots WHERE id = %s",
            (bot_id,)
        )

        conn.commit()

    finally:
        cur.close()
        conn.close()

    # 3️⃣ Delete vector indexes from disk
    delete_faiss_index_for_bot(bot_id)
    delete_csv_faiss_for_bot(bot_id)


def add_file_to_bot(bot_id: str, filename: str, file_type: str, file_size: int):
    """Add a file entry to the bot's file list"""
    conn = _get_db_conn()
    cur = conn.cursor()
    
    file_id = str(uuid.uuid4())
    file_entry = {
        "id": file_id,
        "filename": filename,
        "file_size": file_size,
        "upload_date": time.time(),
        "file_type": file_type
    }
    
    if file_type == "csv":
        cur.execute(
            "UPDATE bots SET csv_files = csv_files || %s WHERE id = %s",
            (json.dumps([file_entry]), bot_id)
        )
    else:  # doc files
        cur.execute(
            "UPDATE bots SET doc_files = doc_files || %s WHERE id = %s",
            (json.dumps([file_entry]), bot_id)
        )
    
    conn.commit()
    cur.close()
    conn.close()
    
    return file_id


def remove_file_from_bot(bot_id: str, file_id: str, file_type: str):
    """Remove a file entry from the bot's file list and its stored content"""
    conn = _get_db_conn()
    cur = conn.cursor()
    
    if file_type == "csv":
        # Remove file from csv_files
        cur.execute(
            "UPDATE bots SET csv_files = (SELECT COALESCE(jsonb_agg(elem), '[]'::jsonb) FROM jsonb_array_elements(csv_files) elem WHERE elem->>'id' != %s) WHERE id = %s",
            (file_id, bot_id)
        )

        # Check remaining CSV files
        cur.execute("SELECT csv_files FROM bots WHERE id = %s", (bot_id,))
        csv_files = cur.fetchone()[0]

        if not csv_files or len(csv_files) == 0:
            # No CSV files remain → clear all csv_content
            cur.execute("UPDATE bots SET csv_content = '[]'::jsonb WHERE id = %s", (bot_id,))
        else:
            # Filter csv_content to keep only rows whose file_id is still present.
            remaining_ids = [f["id"] for f in csv_files if isinstance(f, dict) and "id" in f]
            if remaining_ids:
                cur.execute(
                    """
                    UPDATE bots
                    SET csv_content = (
                        SELECT COALESCE(jsonb_agg(elem), '[]'::jsonb)
                        FROM jsonb_array_elements(csv_content) elem
                        WHERE (elem ? 'file_id' AND elem->>'file_id' = ANY(%s))
                    )
                    WHERE id = %s
                    """,
                    (remaining_ids, bot_id)
                )

    else:  # doc files
        # Remove file from doc_files
        cur.execute(
            "UPDATE bots SET doc_files = (SELECT COALESCE(jsonb_agg(elem), '[]'::jsonb) FROM jsonb_array_elements(doc_files) elem WHERE elem->>'id' != %s) WHERE id = %s",
            (file_id, bot_id)
        )

        # Check remaining doc files
        cur.execute("SELECT doc_files FROM bots WHERE id = %s", (bot_id,))
        doc_files = cur.fetchone()[0]

        if not doc_files or len(doc_files) == 0:
            # No doc files remain → clear all doc_content
            cur.execute("UPDATE bots SET doc_content = '[]'::jsonb WHERE id = %s", (bot_id,))
        else:
            # Filter doc_content to keep only records whose file_id is still present.
            remaining_ids = [f["id"] for f in doc_files if isinstance(f, dict) and "id" in f]
            if remaining_ids:
                cur.execute(
                    """
                    UPDATE bots
                    SET doc_content = (
                        SELECT COALESCE(jsonb_agg(elem), '[]'::jsonb)
                        FROM jsonb_array_elements(doc_content) elem
                        WHERE (elem ? 'file_id' AND elem->>'file_id' = ANY(%s))
                    )
                    WHERE id = %s
                    """,
                    (remaining_ids, bot_id)
                )
    
    conn.commit()
    cur.close()
    conn.close()


def get_bot_files(bot_id: str):
    """Get all files associated with a bot"""
    conn = _get_db_conn()
    cur = conn.cursor()
    
    cur.execute(
        "SELECT csv_files, doc_files FROM bots WHERE id = %s",
        (bot_id,)
    )
    
    row = cur.fetchone()
    cur.close()
    conn.close()
    
    if not row:
        return {"csv_files": [], "doc_files": []}
    
    return {
        "csv_files": row[0] if row[0] else [],
        "doc_files": row[1] if row[1] else []
    }


def clear_bot_files(bot_id: str, file_type: str = None):
    """Clear files for a bot. If file_type is None, clears both CSV and doc files"""
    conn = _get_db_conn()
    cur = conn.cursor()
    
    if file_type == "csv" or file_type is None:
        cur.execute("UPDATE bots SET csv_files = %s WHERE id = %s", (json.dumps([]), bot_id))
        cur.execute("UPDATE bots SET csv_content = '[]'::jsonb WHERE id = %s", (bot_id,))
    
    if file_type == "doc" or file_type is None:
        cur.execute("UPDATE bots SET doc_files = %s WHERE id = %s", (json.dumps([]), bot_id))
        cur.execute("UPDATE bots SET doc_content = '[]'::jsonb WHERE id = %s", (bot_id,))
    
    conn.commit()
    cur.close()
    conn.close()
    

# ---------- Helpers: extract text from uploaded files / url ----------
def _extract_text_from_pdf_file(file_obj: BinaryIO) -> str:
    """Extract full text from a PDF file-like object using pdfplumber."""
    text_parts = []
    try:
        with pdfplumber.open(file_obj) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                page_text = clean_pdf_text(page_text)
                text_parts.append(page_text)
    except Exception:
        # Some pdfplumber failures occur on file-like streams; try fallback
        file_obj.seek(0)
        import pypdf
        reader = pypdf.PdfReader(file_obj)
        for p in reader.pages:
            page_text = clean_pdf_text(p.extract_text() or "")
            text_parts.append(page_text)
    return "\n\n".join(text_parts)


def _extract_text_from_docx_file(file_obj: BinaryIO) -> str:
    """Extract text from a DOCX file-like object using python-docx."""
    # python-docx requires a filename or file-like object
    try:
        doc = docx.Document(file_obj)
    except Exception:
        # fallback: save to temp file and read
        import tempfile
        file_obj.seek(0)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(file_obj.read())
            tmp.flush()
            doc = docx.Document(tmp.name)
    texts = []
    for para in doc.paragraphs:
        texts.append(para.text)
    return "\n\n".join(texts)


def _extract_text_from_txt_file(file_obj: BinaryIO, encoding: str = "utf-8") -> str:
    file_obj.seek(0)
    return file_obj.read().decode(encoding, errors="ignore") if isinstance(file_obj.read(0), bytes) else file_obj.read()


def _fetch_and_extract_text_from_url(url: str) -> str:
    """Fetch URL and extract visible text (simple approach with BeautifulSoup)."""
    resp = requests.get(url, timeout=15)
    resp.raise_for_status()
    ct = resp.headers.get("content-type", "")
    if "pdf" in ct or url.lower().endswith(".pdf"):
        from io import BytesIO
        return _extract_text_from_pdf_file(BytesIO(resp.content))
    html = resp.text
    soup = BeautifulSoup(html, "html.parser")

    # remove scripts/styles
    for s in soup(["script", "style", "noscript", "header", "footer", "svg", "nav"]):
        s.decompose()
    visible_text = " ".join([t.strip() for t in soup.stripped_strings])
    return visible_text


def clean_pdf_text(text: str) -> str:
    """
    Clean noisy PDF text:
    - remove page numbers
    - remove very short junk lines
    - normalize whitespace
    """
    lines = text.splitlines()
    cleaned = []

    for line in lines:
        line = line.strip()

        # remove page numbers like "12"
        if re.fullmatch(r"\d+", line):
            continue

        # skip very short noise
        if len(line) < 3:
            continue

        cleaned.append(line)

    # join lines smoothly
    text = " ".join(cleaned)

    # normalize spaces
    text = re.sub(r"\s+", " ", text)

    return text.strip()


# ---------- Helper: build documents list from uploaded files or text/url ----------
def build_documents_from_uploads(
    *,
    uploaded_files: Optional[List[Tuple[str, BinaryIO]]] = None,
    texts: Optional[List[Dict[str, str]]] = None,
    urls: Optional[List[str]] = None
) -> List[Dict[str, Any]]:
    """
    Construct the `documents` list for index_documents_for_bot.
    - uploaded_files: list of tuples (filename, fileobj)
    - texts: list of {"text": "...", "source": "optional name"}
    - urls: list of urls to fetch and extract
    Returns list of {"text","source","page","meta"}
    """
    docs: List[Dict[str, Any]] = []

    # handle uploaded files
    if uploaded_files:
        for filename, fileobj in uploaded_files:
            fileobj.seek(0)
            lower = filename.lower()
            try:
                if lower.endswith(".pdf"):
                    content = _extract_text_from_pdf_file(fileobj)
                elif lower.endswith(".docx") or lower.endswith(".doc"):
                    content = _extract_text_from_docx_file(fileobj)
                elif lower.endswith(".txt"):
                    fileobj.seek(0)
                    content = fileobj.read().decode("utf-8", errors="ignore") if isinstance(fileobj.read(0), bytes) else fileobj.read()
                else:
                    # generic fallback: try to read as text
                    try:
                        fileobj.seek(0)
                        content = fileobj.read().decode("utf-8", errors="ignore")
                    except Exception:
                        # try HTML extraction
                        fileobj.seek(0)
                        content = _fetch_and_extract_text_from_url(filename) if filename.startswith("http") else ""
                if content and len(content.strip())>0:
                    docs.append({"text": content, "source": filename, "page": None, "meta": {}})
            except Exception as e:
                # skip problematic files but log
                logger.exception("Failed to extract text from upload %s: %s", filename, e)
                continue

    # handle plain texts (already provided)
    if texts:
        for t in texts:
            txt = t.get("text") or ""
            src = t.get("source") or "inline_text"
            docs.append({"text": txt, "source": src, "page": None, "meta": t.get("meta", {})})

    # handle urls
    if urls:
        for u in urls:
            try:
                content = _fetch_and_extract_text_from_url(u)
                docs.append({"text": content, "source": u, "page": None, "meta": {}})
            except Exception:
                logger.exception("Failed to fetch/extract url %s", u)
                continue

    return docs
